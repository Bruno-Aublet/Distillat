"""Export de la fiche de livre (résumé, personnages, analyse) en fichier Word (.docx)."""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.book_report import BookReport

HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
COVER_WIDTH_INCHES = 2.2


def _add_title_page(document: Document, result: BookReport) -> None:
    if result.cover_image:
        try:
            cover_paragraph = document.add_paragraph()
            cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_paragraph.add_run().add_picture(
                io.BytesIO(result.cover_image), width=Inches(COVER_WIDTH_INCHES)
            )
        except Exception:  # noqa: BLE001 - une couverture illisible ne doit pas bloquer l'export
            pass

    title = document.add_heading(result.book_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(result.author)
    run.italic = True
    run.font.size = Pt(14)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run("Fiche de lecture générée automatiquement en français")
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    document.add_paragraph()


def _add_body_text(document: Document, text: str) -> None:
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=2)
        else:
            paragraph = document.add_paragraph(line)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.15


def _add_characters_section(document: Document, result: BookReport) -> None:
    document.add_page_break()
    document.add_heading("Personnages principaux", level=1)

    if not result.characters:
        document.add_paragraph("Aucun personnage principal identifié.")
        return

    for character in result.characters:
        document.add_heading(character.name, level=2)
        paragraph = document.add_paragraph(character.description)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.15


def _add_analysis_section(document: Document, result: BookReport) -> None:
    document.add_page_break()
    document.add_heading("Analyse littéraire", level=1)
    _add_body_text(document, result.analysis_text)


def export_book_report_to_docx(result: BookReport, output_path: str) -> None:
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_title_page(document, result)

    document.add_heading("Résumé", level=1)
    _add_body_text(document, result.summary_text)

    if result.detailed_summary_text:
        document.add_page_break()
        document.add_heading("Résumé détaillé", level=1)
        _add_body_text(document, result.detailed_summary_text)

    _add_characters_section(document, result)
    _add_analysis_section(document, result)

    document.save(output_path)
